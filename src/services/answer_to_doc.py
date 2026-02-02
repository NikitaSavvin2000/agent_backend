import os
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
import base64
import re
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import tempfile
from io import BytesIO

home_path = os.getcwd()

# mock_doc_path = os.path.join(home_path, "src", "mocks", "mock_doc.docx")
# mock_path = os.path.join(home_path, "src", "mocks", "graph.png")
# mock_html_path = os.path.join(home_path, "src", "mocks", "mock_html.html")
# mock_doc_path = os.path.join(home_path, "src", "mocks", "mock_doc.docx")



def remove_emojis(text):
    # Убираем любые символы, которые не буквы, цифры, знаки препинания и пробелы
    return re.sub(r"[^\w\s.,:;!?()\-–—]", "", text)


def add_horizontal_line(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r = p._element
    pPr = p_r.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')       # тип линии: сплошная
    top.set(qn('w:sz'), '6')             # толщина линии
    top.set(qn('w:space'), '1')          # отступ
    top.set(qn('w:color'), '000000')     # цвет линии черный
    pBdr.append(top)
    pPr.append(pBdr)


def remove_emojis(text: str) -> str:
    return re.sub(r"[^\w\s.,:;!?()\-–—]", "", text)

def insert_table(doc, table_json):
    if not table_json:
        return
    table_json_preview = table_json[:20]
    truncated = len(table_json) > 20
    keys = list(table_json_preview[0].keys())
    table = doc.add_table(rows=1, cols=len(keys))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, key in enumerate(keys):
        hdr_cells[i].text = str(key)
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for idx, row_data in enumerate(table_json_preview):
        row_cells = table.add_row().cells
        for i, key in enumerate(keys):
            if truncated and idx == len(table_json_preview) - 1:
                row_cells[i].text = "..."
            else:
                row_cells[i].text = str(row_data.get(key, ""))
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_markdown_to_doc(doc: Document, markdown_text: str):
    html = markdown.markdown(markdown_text)
    soup = BeautifulSoup(html, "html.parser")
    elements = soup.find_all(["h1", "h2", "h3", "p", "li"])

    for i, el in enumerate(elements):
        text = remove_emojis(el.get_text())
        if not text:
            continue

        # Разбиваем по переносам строк и вставляем как несколько параграфов
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue

            if el.name in ["h1", "h2", "h3"]:
                run = doc.add_heading(level=int(el.name[1])).add_run(line)
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                run.font.size = Pt({1: 24, 2: 20, 3: 16}[int(el.name[1])])
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif el.name == "p":
                run = doc.add_paragraph().add_run(line)
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif el.name == "li":
                run = doc.add_paragraph(style="List Bullet").add_run(line)
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Добавляем сплошную линию между блоками
        # if i != len(elements) - 1:
        #     add_horizontal_line(doc)


def html_to_png_blocks(html_code: str, block_selector: str = "div", scale: float = 0.67) -> list:

    with tempfile.NamedTemporaryFile("w", suffix=".html", delete=False, encoding="utf-8") as f:
        f.write(html_code)
        html_path = f.name

    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument(f"--force-device-scale-factor={scale}")

    driver = webdriver.Chrome(options=chrome_options)
    driver.get(f"file://{html_path}")

    blocks = driver.find_elements("css selector", block_selector)
    images = []

    for i, block in enumerate(blocks):
        driver.execute_script("arguments[0].scrollIntoView(true);", block)
        png_bytes = block.screenshot_as_png
        images.append(BytesIO(png_bytes))

    driver.quit()
    return images


def markdown_answer_2_doc(markdown_answer: str, message_html_code: str = None, table_json=None):
    doc = Document()

    # --- 1. Markdown ---
    add_markdown_to_doc(doc, markdown_answer)

    # --- 2. Вставляем таблицу, если есть ---
    insert_table(doc, table_json)

    # --- 3. HTML график через Selenium ---
    if message_html_code:
        try:
            png_streams = html_to_png_blocks(message_html_code, block_selector="body > *")

            for img_stream in png_streams:
                doc.add_paragraph("")
                doc.add_picture(img_stream, width=Inches(6))
        except Exception as e:
            print(f"Ошибка при рендере графика: {e}")

    # --- 4. Конвертируем DOCX в base64 ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode("utf-8")


#
# with open(mock_html_path, "r", encoding="utf-8") as f:
#     message_html_code = f.read()
#
# markdown_answer = "# Im marckdown"
#
# doc_base64 = markdown_answer_2_doc(markdown_answer=markdown_answer, message_html_code=message_html_code, table_json=None)
#
# with open(mock_doc_path, "wb") as f:
#     f.write(base64.b64decode(doc_base64))